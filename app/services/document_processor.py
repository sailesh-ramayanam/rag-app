"""Document processing service for extracting text from various file formats."""

import os
from typing import Tuple
from pathlib import Path

import PyPDF2
import pdfplumber
from docx import Document as DocxDocument


class DocumentProcessor:
    """Handles text extraction from documents."""
    
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        
    def extract_text(self) -> Tuple[str, dict]:
        """
        Extract text from document.
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        suffix = self.file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self._extract_from_pdf()
        elif suffix == ".docx":
            return self._extract_from_docx()
        elif suffix == ".txt":
            return self._extract_from_txt()
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def _extract_from_pdf(self) -> Tuple[str, dict]:
        """Extract text from PDF using pdfplumber for better accuracy."""
        text_parts = []
        metadata = {"page_count": 0, "word_count": 0}
        
        try:
            with pdfplumber.open(self.file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception:
            # Fallback to PyPDF2 if pdfplumber fails
            with open(self.file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                metadata["page_count"] = len(reader.pages)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        metadata["word_count"] = len(full_text.split())
        
        return full_text, metadata
    
    def _extract_from_docx(self) -> Tuple[str, dict]:
        """Extract text from DOCX file."""
        doc = DocxDocument(self.file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        metadata = {
            "page_count": None,  # DOCX doesn't have fixed pages
            "word_count": len(full_text.split()),
        }
        
        return full_text, metadata
    
    def _extract_from_txt(self) -> Tuple[str, dict]:
        """Extract text from plain text file."""
        with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        metadata = {
            "page_count": None,
            "word_count": len(text.split()),
        }
        
        return text, metadata


def get_mime_type(filename: str) -> str:
    """Determine MIME type from filename."""
    suffix = Path(filename).suffix.lower()
    
    mime_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
    }
    
    return mime_map.get(suffix, "application/octet-stream")


def is_supported_file(filename: str) -> bool:
    """Check if file type is supported."""
    suffix = Path(filename).suffix.lower()
    return suffix in [".pdf", ".docx", ".txt"]

